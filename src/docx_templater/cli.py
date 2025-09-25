import pandas as pd
from docx import Document
import re
import os
import argparse


def load_csv_data(csv_path):
    try:
        with open(csv_path, 'r', encoding='utf-8') as f:
            sample = f.read(1024)

        delimiters = [';', ',', '\t', '|']
        delimiter_counts = {delim: sample.count(delim) for delim in delimiters}
        detected_delimiter = max(delimiter_counts, key=delimiter_counts.get)

        data = pd.read_csv(csv_path, delimiter=detected_delimiter)

        data.columns = data.columns.str.strip()
        csv_columns = set(data.columns)

        return data, csv_columns
    except Exception as e:
        print(f"Error loading CSV: {e}")
        return None, set()


def extract_variables_from_docx(doc):
    variables = set()

    def extract_from_paragraphs(paragraphs):
        for paragraph in paragraphs:
            text = paragraph.text
            found_vars = re.findall(r'\{([^}]+)\}', text)
            variables.update(found_vars)

    extract_from_paragraphs(doc.paragraphs)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                extract_from_paragraphs(cell.paragraphs)

    for section in doc.sections:
        if section.header:
            extract_from_paragraphs(section.header.paragraphs)
        if section.footer:
            extract_from_paragraphs(section.footer.paragraphs)

    return variables


def replace_in_paragraph_xml(paragraph, pattern, replacement):
    runs = paragraph.runs
    char_map = [(run_idx, char_idx, char)
                for run_idx, run in enumerate(runs)
                for char_idx, char in enumerate(run.text)]

    combined_text = ''.join(char for _, _, char in char_map)
    pattern_start = combined_text.find(pattern)

    if pattern_start == -1:
        return

    pattern_end = pattern_start + len(pattern)

    affected_runs = {char_map[i][0] for i in range(pattern_start, pattern_end) if i < len(char_map)}
    if not affected_runs:
        return

    for run in runs:
        if pattern in run.text:
            run.text = run.text.replace(pattern, replacement)
            return

    if len(affected_runs) > 1:
        affected_run_indices = sorted(affected_runs)
        combined_affected_text = ''.join(runs[run_idx].text for run_idx in affected_run_indices)
        new_combined_text = combined_affected_text.replace(pattern, replacement, 1)

        for i, run_idx in enumerate(affected_run_indices):
            runs[run_idx].text = new_combined_text if i == 0 else ''


def replace_text_in_paragraph(paragraph, replacements):
    full_text = paragraph.text
    variables_to_replace = [(f"{{{var_name}}}", str(value))
                            for var_name, value in replacements.items()
                            if f"{{{var_name}}}" in full_text]

    for pattern, value in variables_to_replace:
        replace_in_paragraph_xml(paragraph, pattern, value)


def replace_text_in_cell(cell, replacements):
    full_cell_text = cell.text
    if not any(f"{{{var_name}}}" in full_cell_text for var_name in replacements):
        return
    for paragraph in cell.paragraphs:
        replace_text_in_paragraph(paragraph, replacements)


def replace_variables_in_docx(doc, row_data):
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, row_data)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_in_cell(cell, row_data)

    for section in doc.sections:
        if section.header:
            for paragraph in section.header.paragraphs:
                replace_text_in_paragraph(paragraph, row_data)
        if section.footer:
            for paragraph in section.footer.paragraphs:
                replace_text_in_paragraph(paragraph, row_data)


def compare_variables_and_columns(template_variables, csv_columns):
    matched = [var for var in template_variables if var in csv_columns]
    unmatched_template = [var for var in template_variables if var not in csv_columns]
    unmatched_csv = [col for col in csv_columns if col not in template_variables]

    if unmatched_template:
        print("\nWarning: The following placeholders are in the template but missing in the CSV:")
        for v in unmatched_template:
            print(f"   - {v}")

    if unmatched_csv:
        print("\nNote: The following CSV columns are unused in the template:")
        for c in unmatched_csv:
            print(f"   - {c}")

    return {
        "matched": matched,
        "unmatched_template": unmatched_template,
        "unmatched_csv": unmatched_csv
    }


def process_documents(template_path, csv_path, output_dir, filename_pattern):
    os.makedirs(output_dir, exist_ok=True)

    data, csv_columns = load_csv_data(csv_path)
    if data is None:
        return

    try:
        doc = Document(template_path)
        template_vars = extract_variables_from_docx(doc)
    except Exception as e:
        print(f"Error loading template: {e}")
        return

    compare_variables_and_columns(template_vars, csv_columns)

    for index, row in data.iterrows():
        doc = Document(template_path)

        row_data = {}
        for var in template_vars:
            if var in data.columns:
                value = row[var]
                row_data[var] = "" if pd.isna(value) else str(value)

        replace_variables_in_docx(doc, row_data)

        raw_filename = filename_pattern.format(**row_data)
        output_path = os.path.join(output_dir, raw_filename)

        try:
            doc.save(output_path)
            print(f"Created: {output_path}")
        except Exception as e:
            print(f"Error saving {output_path}: {e}")


def main():
    parser = argparse.ArgumentParser(
        description="Generate Word documents from a template and CSV data."
    )
    parser.add_argument(
        "--template", "-t", required=True, help="Path to the Word template (.docx)"
    )
    parser.add_argument(
        "--csv", "-c", required=True, help="Path to the CSV data file"
    )
    parser.add_argument(
        "--outdir", "-o", required=True, help="Directory where output files will be saved"
    )
    parser.add_argument(
        "--pattern", "-p", required=True,
        help="Filename pattern, e.g. '{document_number}_{name}.docx'"
    )

    args = parser.parse_args()

    process_documents(args.template, args.csv, args.outdir, args.pattern)
    print("\nProcessing complete.")


if __name__ == "__main__":
    main()
